import io
from typing import Optional


def export_docx(chapters: list[dict]) -> bytes:
    """导出为 DOCX 格式"""
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()
    doc.add_heading("小说导出", level=0)

    for ch in chapters:
        doc.add_heading(f"第{ch.get('chapterIndex', '')}章 {ch.get('chapter', '')}", level=1)
        for para in (ch.get("chapterData", "") or "").split("\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.style.font.size = Pt(12)
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf(chapters: list[dict]) -> bytes:
    """导出为 PDF 格式"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    styles = getSampleStyleSheet()

    # 尝试注册中文字体
    try:
        pdfmetrics.registerFont(TTFont("SimSun", "SimSun.ttf"))
        body_style = ParagraphStyle("ChineseBody", parent=styles["Normal"], fontName="SimSun", fontSize=12, leading=20)
        title_style = ParagraphStyle("ChineseTitle", parent=styles["Heading1"], fontName="SimSun", fontSize=16, leading=24)
    except Exception:
        body_style = styles["Normal"]
        title_style = styles["Heading1"]

    story = []
    for ch in chapters:
        story.append(Paragraph(f"第{ch.get('chapterIndex', '')}章 {ch.get('chapter', '')}", title_style))
        story.append(Spacer(1, 0.5 * cm))
        for para in (ch.get("chapterData", "") or "").split("\n"):
            if para.strip():
                story.append(Paragraph(para.strip(), body_style))
                story.append(Spacer(1, 0.2 * cm))
        story.append(PageBreak())

    doc.build(story)
    return buf.getvalue()


def export_epub(chapters: list[dict], title: str = "小说导出", author: str = "Toonflow") -> bytes:
    """导出为 EPUB 格式"""
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier("toonflow-export")
    book.set_title(title)
    book.set_language("zh")
    book.add_author(author)

    spine = ["nav"]
    toc = []

    for ch in chapters:
        ch_title = f"第{ch.get('chapterIndex', '')}章 {ch.get('chapter', '')}"
        filename = f"chapter_{ch.get('chapterIndex', 0)}.xhtml"
        content = (ch.get("chapterData", "") or "").replace("\n", "<br/>")

        epub_ch = epub.EpubHtml(title=ch_title, file_name=filename, lang="zh")
        epub_ch.content = f"<h1>{ch_title}</h1><p>{content}</p>"
        book.add_item(epub_ch)
        spine.append(epub_ch)
        toc.append(epub_ch)

    book.toc = toc
    book.spine = spine
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    buf = io.BytesIO()
    epub.write_epub(buf, book)
    return buf.getvalue()
